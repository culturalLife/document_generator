"""
report_generator.py — Core Mistral Report Generation Pipeline
==============================================================
Hardened production version of new_docs_report.py.

Key changes from the original:
  - All config values (model, temperature, parallelism) pulled from config.py
  - print() replaced with structured logger calls
  - fetch_section() wrapped with retry + exponential backoff
  - request_id threaded through all log lines for traceability
  - Per-part timing logged so you can see which call is slow
  - Exceptions are re-raised cleanly so api.py can handle them properly
"""

import os
import re
import time
import concurrent.futures
from pathlib import Path
from typing import Dict

from dotenv import load_dotenv
try:
    from mistralai import Mistral
except ImportError:
    from mistralai.client import Mistral
# pyrefly: ignore [missing-import]
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import config # pyrefly: ignore [missing-import]
from logger import get_logger # pyrefly: ignore [missing-import]

logger = get_logger(__name__)

# ─── Mistral Client Initialization ─────────────────────────────────────────────
# Load environment variables (support local repository root or VM-style local directory)
if (config.BASE_DIR / ".env").exists():
    load_dotenv(config.BASE_DIR / ".env")
else:
    load_dotenv(config.BASE_DIR.parent.parent / ".env")

_api_key = os.environ.get("MISTRAL_API_KEY")
if not _api_key:
    raise EnvironmentError(
        "MISTRAL_API_KEY not found. Set it in the .env file or as an environment variable on the VM."
    )

# Raise the httpx read timeout so long mistral-large-latest calls don't hit the
# default ~60s cutoff. MISTRAL_READ_TIMEOUT can be overridden via the .env file.
_READ_TIMEOUT: float = float(os.environ.get("MISTRAL_READ_TIMEOUT", "400"))
client = Mistral(api_key=_api_key, timeout_ms=int(_READ_TIMEOUT * 1000))
logger.info(
    f"Mistral client initialized | model={config.GENERATION_MODEL} "
    f"| read_timeout={_READ_TIMEOUT:.0f}s"
)


# ─── Retry Helper ──────────────────────────────────────────────────────────────

def _call_with_retry(func, *args, request_id: str = "", part_num: int = 0, **kwargs):
    """
    Calls func(*args, **kwargs) with exponential backoff retry.
    Retries on any exception up to config.MAX_RETRIES times.

    Backoff schedule (default config):
      Attempt 1 fails → wait 2s
      Attempt 2 fails → wait 4s
      Attempt 3 fails → wait 8s → raise
    """
    last_exc = None
    for attempt in range(1, config.MAX_RETRIES + 1):
        try:
            return func(*args, **kwargs)
        except Exception as exc:
            last_exc = exc
            if attempt < config.MAX_RETRIES:
                wait = config.RETRY_BACKOFF_SECONDS * (2 ** (attempt - 1))
                logger.warning(
                    f"[{request_id}] [Part {part_num}] Attempt {attempt}/{config.MAX_RETRIES} failed: "
                    f"{type(exc).__name__}: {exc} — retrying in {wait:.1f}s"
                )
                time.sleep(wait)
            else:
                logger.error(
                    f"[{request_id}] [Part {part_num}] All {config.MAX_RETRIES} attempts exhausted. "
                    f"Final error: {type(exc).__name__}: {exc}"
                )
    raise last_exc


# ─── Section Fetcher ───────────────────────────────────────────────────────────

def _fetch_section(
    part_num: int,
    sections_desc: str,
    prompt_text: str,
    data_summary_text: str,
    request_id: str = "",
) -> str:
    """
    Makes a single Mistral chat.complete() call to generate one section group.
    Called in parallel by generate_report() via ThreadPoolExecutor.
    Wrapped in retry logic.
    """
    t_start = time.monotonic()
    logger.info(f"[{request_id}] [Part {part_num}] Starting API call | model={config.GENERATION_MODEL}")

    system_prompt = config.SYSTEM_PROMPT_TEMPLATE.format(data_summary_text=data_summary_text)

    directive = (
        f"\n\nCRITICAL DIRECTIVE:\n"
        f"You are responsible for generating ONLY the following parts of the report:\n"
        f"{sections_desc}\n"
        f"Do NOT generate any other sections. Begin directly with the content. Do not output any preamble or postamble."
    )

    def _do_call():
        return client.chat.complete(
            model=config.GENERATION_MODEL,
            temperature=config.TEMPERATURE,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt_text + directive},
            ],
        )

    response = _call_with_retry(
        _do_call, request_id=request_id, part_num=part_num
    )

    content = response.choices[0].message.content.strip()
    elapsed = time.monotonic() - t_start
    logger.info(f"[{request_id}] [Part {part_num}] Completed in {elapsed:.1f}s | chars={len(content)}")
    return content


# ─── Document Helpers ──────────────────────────────────────────────────────────

def _apply_bold_markdown(paragraph, text: str):
    """Parses **bold** and *italic* markdown and applies it to a docx paragraph."""
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    lines = text.split('\n')
    for idx, line in enumerate(lines):
        if idx > 0:
            paragraph.add_run().add_break()
        bold_parts = re.split(r'(\*\*.*?\*\*)', line)
        for b_part in bold_parts:
            if b_part.startswith('**') and b_part.endswith('**'):
                run = paragraph.add_run(b_part[2:-2])
                run.bold = True
            else:
                italic_parts = re.split(r'(\*.*?\*)', b_part)
                for i_part in italic_parts:
                    if i_part.startswith('*') and i_part.endswith('*'):
                        run = paragraph.add_run(i_part[1:-1])
                        run.italic = True
                    else:
                        paragraph.add_run(i_part)


def _set_cell_background(cell, fill_color: str):
    """Sets a background color hex string for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)


def _add_horizontal_line(paragraph):
    """Adds a styled horizontal divider using paragraph borders."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B0C4DE')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _clean_heading_text(text: str) -> str:
    """Strips errant '**' from heading text."""
    return text.strip().replace('**', '')


# ─── Document Assembler ────────────────────────────────────────────────────────

def _assemble_docx(content: str, output_path: str, request_id: str = ""):
    """Parses markdown-flavoured LLM output and writes a styled boardroom .docx."""
    logger.info(f"[{request_id}] Assembling Word document...")

    doc = Document()
    styles = doc.styles

    # Custom styles
    title_style = styles.add_style('ReportTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x0F, 0x20, 0x4B)

    subtitle_style = styles.add_style('ReportSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.name = 'Arial'
    subtitle_style.font.size = Pt(14)
    subtitle_style.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    h1_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    h2_style = styles.add_style('SubSectionHeading', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(13)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)

    # Parse content lines
    lines = content.split('\n')
    in_table = False
    current_table = None

    for line in lines:
        line = line.strip()

        if line.startswith('```'):
            continue
        line = line.replace('`', '')

        if not line:
            if in_table:
                in_table = False
                current_table = None
            continue

        # Tables
        if line.startswith('|') and line.endswith('|'):
            if not in_table:
                in_table = True
                cols = [c.strip() for c in line.split('|')[1:-1]]
                current_table = doc.add_table(rows=1, cols=len(cols))
                current_table.style = 'Table Grid'
                hdr_cells = current_table.rows[0].cells
                for i, col_name in enumerate(cols):
                    hdr_cells[i].paragraphs[0].add_run(_clean_heading_text(col_name)).bold = True
                    _set_cell_background(hdr_cells[i], 'D9D9D9')
            else:
                if set(line.replace('|', '').replace('-', '').replace(':', '').strip()) == set():
                    continue  # separator row
                cols = [c.strip() for c in line.split('|')[1:-1]]
                while len(cols) < len(current_table.columns):
                    cols.append("")
                row_cells = current_table.add_row().cells
                for i, col_text in enumerate(cols[:len(current_table.columns)]):
                    _apply_bold_markdown(row_cells[i].paragraphs[0], col_text)
            continue
        else:
            if in_table:
                in_table = False
                current_table = None

        # Horizontal dividers
        if line in ('---', '***'):
            _add_horizontal_line(doc.add_paragraph())
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = _clean_heading_text(heading_match.group(2))
            if level <= 2:
                p = doc.add_paragraph(heading_text, style='SectionHeading')
                _add_horizontal_line(p)
            else:
                doc.add_paragraph(heading_text, style='SubSectionHeading')
            continue

        # Bullet lists
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            bullet_text = line[2:].strip()
            if bullet_text.startswith('**') and bullet_text.endswith('**'):
                bullet_text = bullet_text[2:-2].strip()
            _apply_bold_markdown(p, bullet_text)
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+\.\s)(.*)', line)
        if num_match:
            p = doc.add_paragraph(style='List Number')
            num_text = num_match.group(2).strip()
            if num_text.startswith('**') and num_text.endswith('**'):
                num_text = num_text[2:-2].strip()
            _apply_bold_markdown(p, num_text)
            continue

        # Normal paragraph
        _apply_bold_markdown(doc.add_paragraph(), line)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info(f"[{request_id}] Document saved → {output_path}")


# ─── Section Splitter ─────────────────────────────────────────────────────────

# Matches lines that look like report section headings (after stripping
# leading whitespace per line):
#   ## Title
#   ### 1. Title
#   ### Title
#   1. Title   (top-level numbered list items)
_SECTION_HEADING_RE = re.compile(
    r"^(?:#{2,4}\s+(?:\d+\.?\s*)?|\d+\.\s)(.+)",
)


def _extract_sections(prompt_text: str) -> list:
    """
    Scans the user's prompt for section headings and returns them as an
    ordered list of title strings.

    Strips leading whitespace from every line before matching so that
    indented headings (e.g. from raw.txt triple-quoted strings) are detected.

    Detection priority (first-match per line):
      1. Markdown headings:  ## Title / ### 1. Title / #### Title
      2. Numbered list items at the start of a line: 1. Title

    Returns an empty list if no headings are found (triggers single-call fallback).
    """
    sections = []
    for line in prompt_text.splitlines():
        stripped = line.strip()
        m = _SECTION_HEADING_RE.match(stripped)
        if m:
            sections.append(m.group(1).strip())
    return sections


def _build_parts(sections: list, max_parallel: int, min_per_call: int = 1) -> list:
    """
    Splits a flat list of section titles into (part_num, sections_desc) tuples
    ready for ThreadPoolExecutor.

    Rules:
      - Never creates more workers than there are sections.
      - Respects min_per_call so tiny prompts don't over-split.
      - Last chunk absorbs any remainder sections.

    Args:
        sections:     Ordered list of section title strings.
        max_parallel: Upper bound on parallel workers (from config).
        min_per_call: Minimum sections per worker (from config).

    Returns:
        List of (part_num, directive_string) tuples.
    """
    n = min(max_parallel, max(1, len(sections) // max(min_per_call, 1)))
    n = max(1, min(n, len(sections)))  # clamp: 1 <= n <= len(sections)

    chunk_size = len(sections) // n
    parts = []
    for i in range(n):
        start = i * chunk_size
        # Last chunk takes everything that remains
        end = start + chunk_size if i < n - 1 else len(sections)
        chunk = sections[start:end]
        bullet_list = "\n".join(f"  - {s}" for s in chunk)
        desc = (
            f"Generate ONLY the following sections of the report (in this exact order):\n"
            f"{bullet_list}\n"
            f"Do NOT generate any other sections. "
            f"Start directly with the first section heading. No preamble or postamble."
        )
        parts.append((i + 1, desc))

    return parts


# ─── Public Entry Point ────────────────────────────────────────────────────────

def generate_report(
    prompt_text: str,
    data_summary_text: str,
    output_path: str,
    request_id: str = "",
) -> str:
    """
    Generates the report via parallel Mistral API calls and saves it as a .docx.

    Args:
        prompt_text:       The full prompt (extracted from combined_input).
        data_summary_text: Summarized JSON payload (from summarize_data.clean_data_summary).
        output_path:       Absolute path where the .docx should be saved.
        request_id:        Short UUID for log correlation (injected by api.py).

    Returns:
        output_path as a string.
    """
    t_total_start = time.monotonic()
    logger.info(
        f"[{request_id}] Report generation started | model={config.GENERATION_MODEL} "
        f"| parallel_calls={config.MAX_PARALLEL_CALLS} "
        f"| section_split={config.SECTION_SPLIT_ENABLED}"
    )

    # ── Determine parallel work items ──────────────────────────────────────────
    if config.SECTION_SPLIT_ENABLED:
        sections = _extract_sections(prompt_text)
        logger.info(f"[{request_id}] Detected {len(sections)} section(s) in prompt")

        if sections:
            parts = _build_parts(
                sections,
                max_parallel=config.MAX_PARALLEL_CALLS,
                min_per_call=config.MIN_SECTIONS_PER_CALL,
            )
            logger.info(
                f"[{request_id}] Split into {len(parts)} parallel part(s) "
                f"(~{len(sections) // len(parts)} sections each)"
            )
        else:
            logger.warning(
                f"[{request_id}] No section headings detected in prompt — "
                f"falling back to single call."
            )
            parts = [
                (
                    1,
                    "Generate the complete report as instructed by the user prompt. "
                    "Include all sections, analysis, tables, and recommendations as specified.",
                )
            ]
    else:
        logger.info(f"[{request_id}] SECTION_SPLIT_ENABLED=False — using single call.")
        parts = [
            (
                1,
                "Generate the complete report as instructed by the user prompt. "
                "Include all sections, analysis, tables, and recommendations as specified.",
            )
        ]

    results: Dict[int, str] = {}

    with concurrent.futures.ThreadPoolExecutor(max_workers=config.MAX_PARALLEL_CALLS) as executor:
        future_map = {
            executor.submit(
                _fetch_section, part_num, desc, prompt_text, data_summary_text, request_id
            ): part_num
            for part_num, desc in parts
        }

        for future in concurrent.futures.as_completed(future_map):
            part_num = future_map[future]
            try:
                results[part_num] = future.result()
            except Exception as exc:
                logger.error(f"[{request_id}] [Part {part_num}] Fatal failure: {exc}", exc_info=True)
                raise RuntimeError(f"Part {part_num} generation failed: {exc}") from exc

    # Stitch all parts in order
    stitched_content = "\n\n".join(results[p] for p in sorted(results.keys()))
    logger.info(f"[{request_id}] All parts received and stitched | total_chars={len(stitched_content)}")

    _assemble_docx(stitched_content, output_path, request_id=request_id)

    total_elapsed = time.monotonic() - t_total_start
    logger.info(f"[{request_id}] Pipeline complete in {total_elapsed:.1f}s → {output_path}")
    return output_path
